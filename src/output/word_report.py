"""Word report generation for methodology, findings, assumptions, and limits."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from src.models import LOW_SOURCE_COVERAGE_FIELDS

logger = logging.getLogger("mobile_de.word")


def generate_word_report(
    path: Path,
    df_vendors: pd.DataFrame,
    df_cars: pd.DataFrame,
    dashboard: dict[str, pd.DataFrame],
    state: str,
    errors: list[dict] | None = None,
    *,
    run_summary: dict | None = None,
    vendor_coverage: pd.DataFrame | None = None,
    vehicle_coverage: pd.DataFrame | None = None,
) -> None:
    """Generate the required Word report document."""
    logger.info("Generating Word report: %s", path)
    path.parent.mkdir(parents=True, exist_ok=True)
    errors = errors or []
    run_summary = run_summary or {}

    doc = Document()
    title = doc.add_heading("Alphabots GmbH - Mobile.de Data Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"State: {state.replace('-', ' ').title()}\n"
        f"Date: {datetime.now().strftime('%d.%m.%Y')}\n"
        "Prepared by: RPA Developer"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    _section(doc, "1. Project Overview")
    doc.add_paragraph(
        "This project is a modular Python scraping and data processing pipeline for the "
        "mobile.de regional dealer directory. It collects dealer contact data and vehicle "
        "listing data, saves raw intermediate files, cleans and classifies vehicles, and "
        "generates dashboard-ready Excel and Word deliverables."
    )

    _section(doc, "2. Data Source and Scraping Scope")
    doc.add_paragraph(
        f"Start URL: https://home.mobile.de/regional/{state}/0.html\n"
        f"Assigned state: {state.replace('-', ' ').title()}\n"
        f"Vendors in this run: {len(df_vendors)}\n"
        f"Vehicles in this run: {len(df_cars)}"
    )

    _section(doc, "3. Extraction Methodology")
    steps = [
        "Use Playwright or curl_cffi through a standard fetch abstraction. Static HTML is attempted first where safe; browser rendering remains available for JavaScript-rendered pages.",
        "Accept the mobile.de cookie consent dialog when it appears.",
        "Paginate the regional state directory and collect dealer homepage links.",
        "Deduplicate vendors by normalized mobile.de dealer URL and assign stable C0000001-style IDs after alphabetic URL sorting.",
        "Visit dealer pages and extract structured JSON-LD and Next.js payload data where available.",
        "Open contact, Über uns, and Impressum sections where present and reveal phone numbers only through visible site controls.",
        "Traverse the known mobile.de vehicle categories for each dealer, including Pkw, Motorräder, Wohnmobile, Lkw, Sattelzugmaschinen, Auflieger, Anhänger, Baumaschinen, Busse, Agrarfahrzeuge, and Stapler.",
        "Collect vehicle records from rendered listing cards and structured searchResults/listing payloads, then paginate dealer inventory pages.",
        "Visit individual vehicle detail pages according to the configured detail policy and extract technical, price, and financing fields when present.",
        "Persist checkpoints and optional SQLite state so interrupted runs can resume without duplicating completed work.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    _section(doc, "4. Headless and Cloud Execution")
    doc.add_paragraph(
        "The scraper supports local headed execution, Firefox/Chrome selection, and Docker/Xvfb server-compatible "
        "execution. Strict headless is technically supported but currently blocked by mobile.de in this environment. "
        "Docker/Xvfb is the recommended server-compatible execution mode. If a page is unavailable or returns access "
        "denied, the failure is recorded with status/debug metadata and reflected in coverage metrics."
    )

    _section(doc, "5. Fields Collected")
    _bullets(
        doc,
        [
            "Vendor fields: Händler ID, Händlername, Standort, PLZ, Städte, Bundesland, Land, telephone numbers, fax, email, homepage, mobile.de link, and total vehicle count.",
            "Vehicle columns and extractors cover Händler ID, Händlername, PLZ, Markes, Models, Fahrzeugtyp, Zustand, Erstzulassung, Kilometerstand, Kraftstoffart, CO2, Preis, Leistung, seats, gearbox, emissions class, color, series, trim, displacement, doors, owners, and financing fields; source coverage is measured per run.",
        ],
    )

    _section(doc, "6. Preprocessing Methodology")
    _bullets(
        doc,
        [
            "Whitespace is trimmed and Unicode text is normalized without transliterating German umlauts.",
            "Prices are parsed to EUR, mileage to km, CO2 to g/km, power to kW and PS, displacement to cm3, durations to months, and interest rates to numeric percentages.",
            "German-formatted prices, mileage, CO2, power, displacement, percentages, first registration, and financing durations are parsed into numeric columns while raw human-readable columns are kept.",
            "Missing fields remain empty; no values are invented or inferred beyond the documented numeric parsing and classifications.",
        ],
    )

    _section(doc, "7. Vehicle Category Classification Rules")
    doc.add_paragraph(
        "Vehicle type is mapped to PKW, Motorrad, Freizeitfahrzeuge, LKW, or Andere using the task mapping. "
        "The final category is accompanied by raw_vehicle_type, normalized_vehicle_type, confidence, and rule metadata. "
        "Unknown or unmapped values are assigned to Andere."
    )

    _section(doc, "8. Manufacturer Origin Classification Rules")
    doc.add_paragraph(
        "Manufacturer grouping follows the task-defined categories, not historical/legal corporate-origin definitions. "
        "For example, Ford and Volvo are assigned to Deutschland because they are listed under Deutschland in the task description. "
        "The final origin is accompanied by raw_manufacturer, normalized_manufacturer, confidence, and rule metadata."
    )

    _section(doc, "9. Data Completeness / Coverage")
    doc.add_paragraph(
        "The project guarantees schema completeness and processing traceability. Source completeness is measured through coverage metrics."
    )
    doc.add_paragraph("Unavailable source values are not guessed.")
    doc.add_paragraph(
        "Schema completeness is guaranteed; source completeness is measured."
    )
    if str(run_summary.get("detail_open_strategy", "")).lower() == "uc-popup":
        doc.add_paragraph(
            "The optional uc-popup detail enrichment opens live listing links in a popup/new tab and can extract "
            "CO₂-Emissionen, Baureihe, Ausstattungslinie, and Anzahl der Fahrzeughalter from real detail pages. "
            "It is slower than listing-payload extraction and is intended for missing-detail enrichment with "
            "low concurrency. Absent source values remain empty and are measured in Data_Coverage and "
            "Requirements_Compliance."
        )
    else:
        doc.add_paragraph(
            "Some vehicle detail fields such as CO₂-Emissionen, Baureihe, Ausstattungslinie, and Anzahl der "
            "Fahrzeughalter can remain low coverage when listing payloads or reachable detail pages do not expose "
            "values. These values are not guessed and are reported transparently in Data_Coverage and "
            "Requirements_Compliance."
        )
    _add_coverage_summary(doc, vendor_coverage, vehicle_coverage)
    _add_low_coverage_fields(doc, vehicle_coverage)
    _add_source_audit_summary(doc, run_summary)

    _section(doc, "10. Dashboard Metrics Explanation")
    _bullets(
        doc,
        [
            "Top and least vendors use the dealer total vehicle count when available and the scraped sample count as a fallback.",
            "Best deal candidates use lower normalized price, lower mileage, newer first registration, better price/kW, and lower CO2 when available.",
            "Worst deal candidates use the inverse ranking of the same deal score.",
            "Efficient vehicles use low CO2, low mileage, reasonable price, and newer first registration. If fuel consumption is unavailable, CO2 and price/km-style indicators are used instead.",
        ],
    )

    _section(doc, "11. Summary of Key Findings")
    _add_key_findings(doc, dashboard)

    _section(doc, "12. Limitations")
    limitations = [
        "The scraper uses responsible browser automation and does not require authentication or private endpoints.",
        "If a page is unavailable or returns access denied, the failure is recorded and reflected in coverage metrics.",
        "For production-grade contractual data access, an authorized API/data partnership would be preferable. This assignment focuses on publicly visible website extraction as requested.",
        "Vehicle detail pages may return temporary 5xx/site-protection responses. After repeated failures, the scraper can continue from structured listing payloads to preserve progress.",
        "Listing payload extractors exist for detail-oriented fields, but CO2-Emissionen, Baureihe, Ausstattungslinie, and Anzahl der Fahrzeughalter remain schema-only when the source payload and detail page do not expose values.",
        "Dealer phone numbers and email addresses may be missing if the dealer does not publish them or if a reveal/contact section cannot be opened.",
        "Financing fields are only populated when a listing exposes financing information.",
        "Vehicle detail page layouts vary by vehicle category, so parsers use multiple robust fallbacks but still leave absent fields empty.",
        "A small test run with vendor/car limits is not representative of the full Nordrhein-Westfalen market.",
    ]
    if str(run_summary.get("detail_open_strategy", "")).lower() == "uc-popup":
        limitations.append(
            "uc-popup detail enrichment is slower than listing-payload extraction and is recommended for targeted missing-field enrichment with concurrency 1."
        )
    if not df_vendors.empty or not df_cars.empty:
        limitations.append(
            "The workbook reflects the data successfully collected in this run, not a guaranteed complete live market census."
        )
    if errors:
        limitations.append(
            f"This run recorded {len(errors)} scraping/runtime errors; see errors.csv/json in the output folder."
        )
    _bullets(doc, limitations)

    _section(doc, "13. Assumptions")
    _bullets(
        doc,
        [
            "A normalized mobile.de dealer homepage URL uniquely identifies a vendor for deduplication.",
            "All vehicles collected from a dealer inventory page belong to that dealer.",
            "The Kategorie/Fahrzeugtyp label is the correct source for the requested vehicle category classification.",
            "German numeric formatting is used on mobile.de pages.",
            "External homepage, phone, fax, and email fields are copied only when visible or present in structured page data.",
        ],
    )

    _section(doc, "14. Output File Descriptions")
    _bullets(
        doc,
        [
            "data/raw/vendors_raw.csv and vendors_raw.json: raw vendor records after dealer scraping.",
            "data/raw/cars_raw.csv and cars_raw.json: raw vehicle records after vehicle detail scraping.",
            "data/processed/cars_processed.csv: cleaned and classified vehicle data.",
            "data/output/mobile_de_nrw_dashboard.xlsx: required workbook with raw, processed, run summary, data coverage, requirements compliance, error, ranking, and dashboard sheets.",
            "data/output/mobile_de_nrw_report.docx: this methodology and findings report.",
        ],
    )

    if run_summary:
        _section(doc, "15. Run Summary")
        summary_keys = [
            "run_id",
            "started_at_utc",
            "finished_at_utc",
            "duration_seconds",
            "target_state",
            "target_state_slug",
            "search_state",
            "pipeline_mode",
            "browser_mode",
            "docker_xvfb_supported",
            "strict_headless_blocked",
            "detail_open_strategy",
            "source_audit_completed",
            "source_audit_dir",
            "source_audit_summary_path",
            "detail_strategy_status",
            "detail_strategy_matrix_path",
            "detail_strategies_tested",
            "detail_strategy_recommendation",
            "source_audit_target_fields_found",
            "processed_vendor_count",
            "extracted_vehicle_count",
            "error_count",
            "vehicle_detail_jobs_total",
            "detail_needed_count",
            "detail_skipped_count",
            "detail_attempted_count",
            "detail_success_count",
            "detail_failed_count",
            "listing_fallback_used_count",
            "detail_page_loaded_count",
            "real_detail_page_loaded_count",
            "detail_target_fields_extracted_count",
            "detail_real_page_but_no_target_fields_count",
            "popup_opened_count",
            "popup_captured_count",
            "popup_capture_failed_count",
            "wrong_tab_capture_count",
            "stale_redirect_count",
            "detail_home_redirect_count",
            "detail_error_page_count",
            "uc_popup_eligible_count",
            "uc_popup_skipped_not_needed_count",
            "uc_popup_skipped_not_applicable_count",
            "uc_popup_attempted_count",
            "uc_popup_success_count",
            "uc_popup_failed_count",
            "uc_popup_avg_seconds",
            "fields_added_by_uc_popup_count",
            "avg_fields_added_per_detail_success",
            "detail_fetch_403_count",
            "detail_fetch_503_count",
            "detail_fetch_failed_count",
            "detail_site_blocked_or_503_count",
            "vendor_required_fields_coverage_pct",
            "vehicle_basic_fields_coverage_pct",
            "vehicle_technical_fields_coverage_pct",
            "financing_fields_coverage_pct",
            "classification_fields_coverage_pct",
            "final_required_fields_coverage_pct",
            "cookie_modal_visible_count",
            "cookie_consent_click_count",
            "cookie_modal_remaining_count",
            "playwright_browser_opened_count",
            "regional_browser_opened_count",
            "vendor_browser_opened_count",
            "vehicle_detail_browser_opened_count",
            "idle_about_blank_count",
        ]
        for key in summary_keys:
            if key in run_summary:
                doc.add_paragraph(f"{key}: {run_summary[key]}")

    doc.save(str(path))
    logger.info("Word report saved: %s", path)


def _section(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_key_findings(doc: Document, dashboard: dict[str, pd.DataFrame]) -> None:
    vendor_summary = dashboard.get("vendor_summary", pd.DataFrame())
    manufacturer_summary = dashboard.get("manufacturer_summary", pd.DataFrame())
    category_summary = dashboard.get("category_summary", pd.DataFrame())
    best_deals = dashboard.get("best_deals", pd.DataFrame())
    efficient = dashboard.get("efficient_vehicles", pd.DataFrame())

    if vendor_summary.empty and manufacturer_summary.empty and category_summary.empty:
        doc.add_paragraph(
            "No market findings could be computed because no usable live records were collected in this run."
        )
        return

    if not vendor_summary.empty:
        top = vendor_summary.iloc[0]
        low = vendor_summary.sort_values(
            ["Total_Vehicle_Count", "Händlername"], ascending=[True, True]
        ).iloc[0]
        doc.add_paragraph(
            f"Top vendor by total vehicle count: {top.get('Händlername', '')} "
            f"({int(top.get('Total_Vehicle_Count', 0))} vehicles)."
        )
        doc.add_paragraph(
            f"Least vendor by total vehicle count: {low.get('Händlername', '')} "
            f"({int(low.get('Total_Vehicle_Count', 0))} vehicles)."
        )

    if not manufacturer_summary.empty:
        high = manufacturer_summary.iloc[0]
        low = manufacturer_summary.iloc[-1]
        doc.add_paragraph(
            f"Highest listed manufacturer: {high.get('Manufacturer', '')} "
            f"({int(high.get('Count', 0))} listings)."
        )
        doc.add_paragraph(
            f"Lowest listed manufacturer: {low.get('Manufacturer', '')} "
            f"({int(low.get('Count', 0))} listings)."
        )

    if not category_summary.empty:
        top_category = category_summary.iloc[0]
        doc.add_paragraph(
            f"Most listed vehicle category: {top_category.get('Category', '')} "
            f"({int(top_category.get('Count', 0))} listings)."
        )

    if not best_deals.empty:
        row = best_deals.iloc[0]
        doc.add_paragraph(
            f"Best deal candidate: {row.get('Markes', '')} {row.get('Models', '')} "
            f"from {row.get('Händlername', '')}."
        )

    if not efficient.empty:
        row = efficient.iloc[0]
        doc.add_paragraph(
            f"Top efficient vehicle candidate: {row.get('Markes', '')} {row.get('Models', '')} "
            f"with CO2 value {row.get('CO₂-Emissionen', '')}."
        )


def _add_coverage_summary(
    doc: Document,
    vendor_coverage: pd.DataFrame | None,
    vehicle_coverage: pd.DataFrame | None,
) -> None:
    rows: list[str] = []
    for label, coverage in [("Vendor", vendor_coverage), ("Vehicle", vehicle_coverage)]:
        if coverage is None or coverage.empty or "coverage_pct" not in coverage.columns:
            continue
        average = pd.to_numeric(coverage["coverage_pct"], errors="coerce").mean()
        rows.append(f"{label} field coverage average: {average:.1f}%.")
    if rows:
        _bullets(doc, rows)
    else:
        doc.add_paragraph("No coverage metrics were available for this run.")


def _add_low_coverage_fields(
    doc: Document, vehicle_coverage: pd.DataFrame | None
) -> None:
    if (
        vehicle_coverage is None
        or vehicle_coverage.empty
        or "field" not in vehicle_coverage.columns
    ):
        return
    coverage = vehicle_coverage.copy()
    coverage["coverage_pct"] = pd.to_numeric(
        coverage.get("coverage_pct"), errors="coerce"
    ).fillna(0.0)
    low_source_rows = coverage[
        coverage["field"].astype(str).isin(LOW_SOURCE_COVERAGE_FIELDS)
        & (coverage["coverage_pct"] < 25)
    ]
    if low_source_rows.empty:
        return
    fields = ", ".join(
        f"{row['field']} ({row['coverage_pct']:.1f}%)"
        for _, row in low_source_rows.sort_values("field").iterrows()
    )
    doc.add_paragraph(f"Low-coverage vehicle detail field group: {fields}.")


def _add_source_audit_summary(doc: Document, run_summary: dict | None) -> None:
    if not run_summary:
        return
    completed = str(run_summary.get("source_audit_completed", "")).lower() == "true"
    if not completed and not run_summary.get("detail_strategy_status"):
        return
    _section(doc, "Source Audit and Detail Strategy Matrix")
    doc.add_paragraph(
        "The source audit checks returned HTML, Next.js payloads, listing-card JSON, visible card text, "
        "network responses, and detail navigation attempts before accepting detail fields as source-limited."
    )
    doc.add_paragraph(f"source_audit_completed: {completed}")
    if run_summary.get("source_audit_dir"):
        doc.add_paragraph(f"source_audit_dir: {run_summary['source_audit_dir']}")
    if run_summary.get("detail_strategies_tested"):
        doc.add_paragraph(
            f"detail_strategies_tested: {run_summary['detail_strategies_tested']}"
        )
    if run_summary.get("detail_strategy_matrix_path"):
        doc.add_paragraph(
            f"detail_strategy_matrix_path: {run_summary['detail_strategy_matrix_path']}"
        )
    if run_summary.get("detail_strategy_recommendation"):
        doc.add_paragraph(
            f"detail_strategy_recommendation: {run_summary['detail_strategy_recommendation']}"
        )
    doc.add_paragraph(
        "Remaining low-coverage detail-dependent fields are CO2-Emissionen, Baureihe, Ausstattungslinie, "
        "and Anzahl der Fahrzeughalter unless a tested detail strategy extracts real values from returned source."
    )
